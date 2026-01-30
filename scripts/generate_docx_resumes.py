import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def set_cell_background(cell, color_hex):
    """
    Set background color for a table cell.
    color_hex: string like "E0E0E0"
    """
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_resume_docx(filename, profile_data, is_roche=False):
    doc = docx.Document()
    
    # Adjust margins to be narrow
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # --- Header ---
    header_table = doc.add_table(rows=1, cols=1)
    header_table.width = Inches(7.5)
    header_cell = header_table.cell(0, 0)
    
    # Name
    name_para = header_cell.paragraphs[0]
    name_run = name_para.add_run("ATUL SINGH")
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_run.font.color.rgb = RGBColor(30, 41, 59) # Dark Slate
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Title
    title_para = header_cell.add_paragraph()
    title_text = "Senior Product Owner" if is_roche else "Senior Product Owner & Strategy Consultant"
    title_run = title_para.add_run(title_text)
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = RGBColor(37, 99, 235) # Blue
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Summary
    summary_para = header_cell.add_paragraph()
    summary_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_run = summary_para.add_run(profile_data['summary'])
    summary_run.italic = True
    summary_run.font.size = Pt(10)
    
    doc.add_paragraph() # Spacer

    # --- Layout Table (Left Sidebar / Right Main) ---
    # We use a table with 1 row, 2 columns
    table = doc.add_table(rows=1, cols=2)
    table.width = Inches(7.5)
    table.autofit = False
    
    # Set Column Widths (approx 30% / 70%)
    # Note: python-docx column width setting is tricky, better to set cell widths if possible
    # But cell.width works best.
    
    left_cell = table.cell(0, 0)
    left_cell.width = Inches(2.2)
    
    right_cell = table.cell(0, 1)
    right_cell.width = Inches(5.3)
    
    # Style styling
    set_cell_background(left_cell, "F8FAFC") # Very light gray for sidebar
    
    # --- FILL LEFT COLUMN (Sidebar) ---
    lc = left_cell.paragraphs[0]
    
    def add_sidebar_header(text, cell):
        p = cell.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(71, 85, 105) # Slate 600
        p_format = p.paragraph_format
        p_format.space_before = Pt(12)
        p_format.space_after = Pt(6)
        
    def add_sidebar_text(text, cell, bold=False, icon=""):
        p = cell.add_paragraph()
        if icon:
            ir = p.add_run(icon + " ")
            ir.font.size = Pt(9)
        r = p.add_run(text)
        if bold: r.bold = True
        r.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(3)

    # Contact
    add_sidebar_header("CONTACT", left_cell)
    add_sidebar_text("Vienna, Austria (Target)", left_cell, icon="📍")
    add_sidebar_text("atulsingh779@gmail.com", left_cell, icon="📧")
    add_sidebar_text("+91 7391-070-594", left_cell, icon="📱")
    add_sidebar_text("LinkedIn Profile", left_cell, icon="🔗")

    # Education
    add_sidebar_header("EDUCATION", left_cell)
    add_sidebar_text("ISB", left_cell, bold=True)
    add_sidebar_text("MBA Equivalent - Exec", left_cell)
    add_sidebar_text("2018-19", left_cell)
    
    left_cell.add_paragraph() # Spacer
    
    add_sidebar_text("NIT Durgapur", left_cell, bold=True)
    add_sidebar_text("B. Tech (Mech)", left_cell)
    add_sidebar_text("2010-14 | CGPA: 8.1", left_cell)

    # Skills
    add_sidebar_header("SKILLS", left_cell)
    for skill in profile_data['skills']:
        add_sidebar_text("• " + skill, left_cell)

    # --- FILL RIGHT COLUMN (Main Content) ---
    rc = right_cell.paragraphs[0]
    
    def add_section_header(text, cell):
        p = cell.add_paragraph()
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(30, 41, 59)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        # Add bottom border simulation if possible? hard in docx. Just keep it simple.

    def add_job_header(title, company, date, cell):
        p = cell.add_paragraph()
        r1 = p.add_run(title)
        r1.bold = True
        r1.font.size = Pt(11)
        
        r2 = p.add_run(f" - {company}")
        r2.italic = True
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor(100, 116, 139)

        # Tab for date or just new line? Tab is better but hard.
        # Let's just put date on next line small.
        p_date = cell.add_paragraph()
        r3 = p_date.add_run(date)
        r3.font.size = Pt(9)
        r3.bold = True
        r3.font.color.rgb = RGBColor(37, 99, 235)
        p_date.paragraph_format.space_after = Pt(2)
        
    def add_bullet(text, cell):
        p = cell.add_paragraph(style='List Bullet')
        p.add_run(text).font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)

    # Experience Section
    add_section_header("Professional Experience", right_cell)

    # Job 1: Lilly
    add_job_header("Senior Consultant / Product Owner", "Eli Lilly and Company", "2019 - Present", right_cell)
    for bullet in profile_data['lilly_exp']:
        add_bullet(bullet, right_cell)
        
    right_cell.add_paragraph()

    # Job 2: ZS
    add_job_header("Decision Analytics Consultant", "ZS Associates, Pune", "Sep 2014 - 2019", right_cell)
    for bullet in profile_data['zs_exp']:
        add_bullet(bullet, right_cell)

    right_cell.add_paragraph()

    # Projects Section
    add_section_header("Key Projects", right_cell)
    add_job_header("Algo Trading Platform", "Private Venture", "Present", right_cell)
    for bullet in profile_data['projects']:
        add_bullet(bullet, right_cell)

    doc.save(filename)
    print(f"Saved {filename}")

# --- DATA ---

common_zs_exp = [
    "Leadership: Led cross-geographic teams of 4-10 associates to deliver ~80 projects for 5 Fortune 500 pharmaceutical clients.",
    "Cost Savings: Achieved $75M savings by optimizing salesforce operations (16% headcount reduction) with minimal disruption.",
    "Operational Excellence: Reduced turnaround time by ~10% via Value Stream Mapping (VSM) sessions.",
    "Performance Improvement: Enhanced sales rep adherence rate by ~15% by incorporating on-ground data.",
    "Portfolio Management: Managed analytics for sales forces >1000 reps and portfolios of 8+ products."
]

common_lilly_exp = [
    "Gen AI Innovation: Leading 'Project Ventu', utilizing Generative AI to decode HCP data; currently finalizing product design.",
    "Omnichannel Execution: Defined and executed Omnichannel strategy as Business Unit Lead for Oncology.",
    "Dynamic Targeting: Product Owner for a dual-component targeting system (Frontend UI & Analytical Backend).",
    "Operational Transformation: Built custom Jira environments to professionalize requirements gathering."
]

common_projects = [
    "Developing 'TradeAlgo', a Python-based algorithmic trading system with real-time risk management and backtesting."
]

# GENERAL RESUME DATA
data_general = {
    'summary': "Experienced Product Owner and Strategy Consultant with over 9 years of experience driving digital transformation. Expert in bridging business needs and technical execution, with focus on Gen AI and Omnichannel strategies.",
    'skills': ["Product Ownership", "Gen AI & Innovation", "Omnichannel Strategy", "Agile / Scrum", "Data Analytics", "Use Python", "Stakeholder Mgmt"],
    'lilly_exp': common_lilly_exp,
    'zs_exp': common_zs_exp,
    'projects': common_projects
}

# ROCHE RESUME DATA
data_roche = {
    'summary': "Purpose-driven Product Owner with experience in digital health and regulated environments. Expert in delivering mobile/cloud solutions within SAFe, leveraging Gen AI and Omnichannel strategies to drive patient outcomes.",
    'skills': ["Product Backlog Mgmt", "SAFe / Scaled Agile", "Mobile & Cloud Apps", "Gen AI Innovation", "Jira & Confluence", "Regulatory Compliance", "Cross-functional Leadership"],
    'lilly_exp': common_lilly_exp, # Same exp, just framed differently in summary/skills
    'zs_exp': common_zs_exp,
    'projects': common_projects + ["Demonstrated hands-on technical understanding of complex systems (Algo Trading)."]
}

if __name__ == "__main__":
    create_resume_docx(r"C:\Users\atuls\Startup\TradeAlgo\data\resume_general.docx", data_general, is_roche=False)
    create_resume_docx(r"C:\Users\atuls\Startup\TradeAlgo\data\resume_roche_spo.docx", data_roche, is_roche=True)
