import docx
import sys
import os

def extract_text(path, output_path):
    if not os.path.exists(path):
        print(f"File not found: {path}")
        return
    
    try:
        doc = docx.Document(path)
        full_text = []
        
        # Helper to clean text
        def clean(t):
            return t.strip()

        # Extract paragraphs
        for para in doc.paragraphs:
            t = clean(para.text)
            if t:
                full_text.append(t)
        
        # Extract tables (often used for skills/exp in resumes)
        for table in doc.tables:
            for row in table.rows:
                row_data = [clean(cell.text) for cell in row.cells if clean(cell.text)]
                if row_data:
                    full_text.append(" | ".join(row_data))
                    
        content = '\n'.join(full_text)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
            
        print(f"Successfully wrote {len(content)} characters to {output_path}")

    except Exception as e:
        print(f"Error extracting text: {e}")

if __name__ == "__main__":
    cv_path = r"C:\Users\atuls\Startup\TradeAlgo\Atul_Singh_CV_v2.1.docx"
    display_path = r"C:\Users\atuls\Startup\TradeAlgo\cv_extracted.txt"
    extract_text(cv_path, display_path)
